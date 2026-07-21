import os
import time
import json
import sys
from pathlib import Path
from dotenv import load_dotenv
from groq import Groq
from pydantic import BaseModel
from pypdf import PdfReader
from docx import Document

_ROOT = os.path.abspath(os.path.join(os.path.dirname(__file__), "..", ".."))
load_dotenv(os.path.join(_ROOT, ".env"))

my_api_key = os.getenv("GROQ_API_KEY")
if not my_api_key:
    raise ValueError("GROQ_API_KEY not found in .env file.")

client = Groq(api_key=my_api_key)
MODEL  = "llama-3.3-70b-versatile"

def api_call_with_retry(messages, max_retries=3):
    for attempt in range(max_retries):
        try:
            response = client.chat.completions.create(
                model=MODEL,
                messages=messages,
                response_format={"type": "json_object"},
            )
            return response
        except Exception as e:
            if "429" in str(e) or "rate" in str(e).lower():
                wait_time = 60 * (attempt + 1)
                print(f"    Rate limited. Waiting {wait_time}s...")
                time.sleep(wait_time)
            else:
                raise
    raise Exception("Max retries exceeded")


class JobD(BaseModel):
    role: str
    required_skills: list[str]
    preferred_skills: list[str]
    minimum_experience: float | None
    education_requirements: list[str]
    responsibilities: list[str]

class Experience(BaseModel):
    company: str | None = None
    role: str | None = None
    duration: str | None = None
    description: str | None = None
    skills_used: list[str] = []

class Resume(BaseModel):
    name: str | None = None
    email: str | None = None
    phone: str | None = None
    total_experience_years: float | None = None
    skills: list[str] = []
    experiences: list[Experience] = []
    education: list[str] = []
    projects: list[str] = []
    certifications: list[str] = []

class MatchResult(BaseModel):
    score: float
    details: dict


def parse_job_description(job_description: str) -> JobD:
    jobd_schema = JobD.model_json_schema()
    system_prompt = f"""
You are an expert HR assistant.
Extract structured information from job descriptions.
Return ONLY valid JSON matching this schema:
{jobd_schema}
Do NOT return schema fields like "properties", "title", or "type".
Fill with actual values. Return null for missing numbers, [] for missing lists.
"""
    messages = [
        {"role": "system", "content": system_prompt},
        {"role": "user", "content": f"Analyze this job description:\n\n{job_description}"},
    ]
    response = api_call_with_retry(messages)
    return JobD(**json.loads(response.choices[0].message.content))


def parse_resume(resume_text: str) -> Resume:
    resume_schema = Resume.model_json_schema()
    system_prompt = f"""
You are an expert resume parser.
Extract information by meaning, not just headings.
Treat Experience, Work History, Employment, Internships as the same.
Collect skills from all sections.
Return ONLY valid JSON matching this schema:
{resume_schema}
Return null for missing values, [] for missing lists. Do not invent data.
"""
    messages = [
        {"role": "system", "content": system_prompt},
        {"role": "user", "content": f"Parse this resume:\n\n{resume_text}"},
    ]
    response = api_call_with_retry(messages)
    return Resume(**json.loads(response.choices[0].message.content))


def final_score(job: JobD, resume: Resume) -> MatchResult:
    match_schema = MatchResult.model_json_schema()
    prompt = f"""
You are an expert HR recruiter with 15+ years of experience. Compare the resume with the job description.

JOB DESCRIPTION:
{job.model_dump_json(indent=2)}

CANDIDATE RESUME:
{resume.model_dump_json(indent=2)}

Return JSON matching this schema:
{match_schema}

The "details" dict must include:
  - candidate_name: Full name of the candidate
  - matching_skills: List of skills that directly match job requirements
  - missing_important_skills: List of critical skills the candidate lacks
  - experience_requirement_met: true/false with explanation
  - overall_match_percentage: 0-100 score
  - final_verdict: A detailed 3-4 sentence professional assessment explaining:
    * Why the candidate scored this percentage
    * Key strengths that match the role
    * Critical gaps that need addressing
    * Recommendation: Strong Hire / Consider / Skip with reasoning
  - strengths: Top 3 specific strengths for this role
  - weaknesses: Top 3 areas for improvement
  - interview_focus: What to probe in technical interview
"""
    messages = [{"role": "user", "content": prompt}]
    response = api_call_with_retry(messages)
    return MatchResult(**json.loads(response.choices[0].message.content))


def read_pdf(file_path: Path) -> str:
    reader = PdfReader(file_path)
    return "\n".join(p.extract_text() for p in reader.pages if p.extract_text())

def read_docx(file_path: Path) -> str:
    doc = Document(file_path)
    text = "\n".join(p.text for p in doc.paragraphs if p.text.strip())
    for table in doc.tables:
        for row in table.rows:
            for cell in row.cells:
                if cell.text.strip():
                    text += "\n" + cell.text
    return text

def read_resume_file(file_path: Path) -> str | None:
    if file_path.suffix.lower() == ".pdf":
        return read_pdf(file_path)
    elif file_path.suffix.lower() == ".docx":
        return read_docx(file_path)
    elif file_path.suffix.lower() == ".txt":
        return file_path.read_text(encoding="utf-8")
    return None


def score_bar(score: float) -> str:
    filled = int(score) // 5
    return "#" * filled + "-" * (20 - filled)

def divider(title: str = "", width: int = 60):
    print("\n" + "=" * width)
    if title:
        print(f"  {title}")
        print("=" * width)

def print_candidate(label: str, candidate: dict):
    print(f"\n  {label}: {candidate['name']}")
    print(f"  Score : {candidate['score']}%  [{score_bar(candidate['score'])}]")
    for key, val in candidate["details"].items():
        if key in ("strengths", "weaknesses", "interview_focus") and isinstance(val, list):
            print(f"  {key.replace('_', ' ').title():<30}:")
            for item in val:
                print(f"    - {item}")
        else:
            print(f"  {key.replace('_', ' ').title():<30}: {val}")


JOB_DESCRIPTIONS = {}

def load_job_descriptions():
    job_files = {
        "sde": "job_description_sde.txt",
        "ml": "job_description_ml.txt",
        "fullstack": "job_description_fullstack.txt",
        "devops": "job_description_devops.txt"
    }
    for key, filename in job_files.items():
        filepath = Path(__file__).parent / filename
        if filepath.exists():
            JOB_DESCRIPTIONS[key] = filepath.read_text(encoding="utf-8")

load_job_descriptions()


def main():
    divider("[AI]  AI RESUME EVALUATOR  -  Groq / LLaMA-3.3-70b")

    job_key = sys.argv[1] if len(sys.argv) > 1 else "sde"
    
    if job_key not in JOB_DESCRIPTIONS:
        print(f"\n  [!] Unknown job: {job_key}")
        print(f"  Available: {', '.join(JOB_DESCRIPTIONS.keys())}")
        return

    print(f"\n[1/3] Parsing Job Description...")
    job = parse_job_description(JOB_DESCRIPTIONS[job_key])
    print(f"      Role              : {job.role}")
    print(f"      Min. Experience   : {job.minimum_experience} yrs")
    print(f"      Education         : {', '.join(job.education_requirements)}")

    resume_folder = Path(__file__).parent / "resumes"
    resume_folder.mkdir(exist_ok=True)

    resume_files = [f for f in resume_folder.iterdir() if f.suffix.lower() in (".pdf", ".docx", ".txt")]

    if not resume_files:
        print(f"\n[!] No resumes found in '{resume_folder}'. Add PDF/DOCX files and re-run.")
        return

    print(f"\n[2/3] Found {len(resume_files)} resume(s). Analysing...\n")

    all_results: list[dict] = []

    for file_path in resume_files:
        print(f"  Processing : {file_path.name}")
        resume_text = read_resume_file(file_path)
        if not resume_text or not resume_text.strip():
            print("    -> Could not extract text. Skipping.\n")
            continue

        parsed = parse_resume(resume_text)
        time.sleep(3)

        result = final_score(job, parsed)
        time.sleep(3)

        if result.score <= 1:
            display_score = result.score * 100
        elif result.score <= 10:
            display_score = result.score * 10
        else:
            display_score = result.score
        
        display_score = min(max(display_score, 0), 100)
        print(f"    -> Score : {display_score:.0f}%\n")
        all_results.append({
            "name":    parsed.name or file_path.stem,
            "score":   display_score,
            "details": result.details,
        })

    if not all_results:
        print("[!] No resumes could be evaluated.")
        return

    all_results.sort(key=lambda c: c["score"], reverse=True)

    divider("TOP 2 CANDIDATES")
    for i, candidate in enumerate(all_results[:2], 1):
        print_candidate(f"Rank #{i}", candidate)

    divider("LOWEST 2 CANDIDATES")
    for i, candidate in enumerate(all_results[-2:], 1):
        print_candidate(f"Bottom #{i}", candidate)

    divider()
    print(f"  [3/3] Done! Evaluated {len(all_results)} candidate(s).")
    print("=" * 60 + "\n")


if __name__ == "__main__":
    main()
